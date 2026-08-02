from docx import Document

from app.core.config import Settings
from app.services import ingestion_service


def test_prepare_text_document_cleans_and_chunks_with_overlap(tmp_path, monkeypatch) -> None:
    source = tmp_path / "knowledge.txt"
    source.write_text("Alpha   beta.\n\n" + "gamma delta " * 30, encoding="utf-8")
    monkeypatch.setattr(
        ingestion_service,
        "get_settings",
        lambda: Settings(chunk_size_characters=120, chunk_overlap_characters=15),
    )

    chunks = ingestion_service.prepare_document(source, "txt")

    assert len(chunks) > 1
    assert chunks[0].page_number is None
    assert "Alpha beta." in chunks[0].content
    assert chunks[0].chunk_index == 0
    assert chunks[-1].chunk_index == len(chunks) - 1


def test_prepare_docx_document_extracts_paragraphs(tmp_path) -> None:
    source = tmp_path / "guide.docx"
    document = Document()
    document.add_paragraph("First instruction.")
    document.add_paragraph("Second instruction.")
    document.save(source)

    chunks = ingestion_service.prepare_document(source, "docx")

    assert "First instruction." in chunks[0].content
    assert "Second instruction." in chunks[0].content
